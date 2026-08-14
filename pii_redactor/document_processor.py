"""
Document Processor
===================
Reads a .docx file, applies PII redaction run-by-run (preserving formatting),
and writes the redacted output.

Key design choices:
  - We operate at the *run* level (docx.Run) rather than paragraph level.
    A run is the smallest formatting unit in a .docx. Editing at this level
    means bold/italic/font/color/size are never touched — only the text changes.
  - We DO NOT rebuild paragraphs from scratch (which would lose all XML
    attributes). Instead we call `redact_text()` and stitch replacement runs
    back in if a run's text is split by a PII span.
  - Tables are iterated cell-by-cell, row-by-row.
  - Headers and footers are iterated through `section.header` / `section.footer`.

Performance note:
  spaCy's NER is called once per paragraph/cell text (not per run) to avoid
  tokenisation overhead. The resulting PIIMatch offsets are then mapped back
  to individual runs via character-offset arithmetic.
"""

from __future__ import annotations
import copy
import logging
import re
from docx import Document
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .detectors import DETECTOR_REGISTRY
from .detectors.base import PIIMatch
from .faker_mapper import FakeValueMapper

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Core text-level redaction
# ---------------------------------------------------------------------------

def _run_detectors(text: str) -> list[PIIMatch]:
    """Run all registered detectors on *text*, return de-duplicated matches."""
    all_matches: list[PIIMatch] = []
    for detector in DETECTOR_REGISTRY:
        try:
            all_matches.extend(detector.detect(text))
        except Exception as exc:
            logger.warning(f"Detector {detector.PII_TYPE} failed: {exc}")

    # Sort by start position, then de-duplicate overlapping spans
    all_matches.sort(key=lambda m: (m.start, -m.confidence))
    deduped: list[PIIMatch] = []
    last_end = -1
    for m in all_matches:
        if m.start >= last_end:
            deduped.append(m)
            last_end = m.end
    return deduped


# Fast pre-filter regex patterns — cheap checks before spaCy NER is called
_RE_HAS_CAPS    = re.compile(r'[A-Z]')          # any uppercase letter
_RE_HAS_DIGIT   = re.compile(r'\d')             # any digit (phones, CINs, dates)
_RE_HAS_AT      = re.compile(r'@')              # email indicator
_RE_HAS_PLUS    = re.compile(r'\+')             # phone prefix

def _might_contain_pii(text: str) -> bool:
    """
    Cheap O(n) pre-filter: returns False for text blocks that cannot
    contain any of our PII types, so we skip the expensive spaCy NER call.

    Rules (all must fail to return False):
      - Blank or very short text (< 4 chars)
      - No uppercase letters  → cannot be a name or CIN
      - No digits             → cannot be phone/email/date/CIN
      - No '@' or '+' signs   → cannot be email or phone
    """
    stripped = text.strip()
    if len(stripped) < 4:
        return False
    # If it has an uppercase letter OR a digit OR @ OR + it might have PII
    return bool(
        _RE_HAS_CAPS.search(stripped)
        or _RE_HAS_DIGIT.search(stripped)
        or _RE_HAS_AT.search(stripped)
        or _RE_HAS_PLUS.search(stripped)
    )


def redact_text(text: str, mapper: FakeValueMapper) -> tuple[str, list[PIIMatch]]:
    """
    Replace PII in *text* with consistent fakes from *mapper*.
    Returns (redacted_text, list_of_matches_found).
    """
    matches = _run_detectors(text)
    if not matches:
        return text, []

    result = []
    prev = 0
    for m in matches:
        result.append(text[prev:m.start])
        fake = mapper.get_or_create(m.pii_type, m.text)
        result.append(fake)
        prev = m.end
    result.append(text[prev:])
    return "".join(result), matches


# ---------------------------------------------------------------------------
# Run-level redaction (preserves formatting)
# ---------------------------------------------------------------------------

def _redact_paragraph(paragraph, mapper: FakeValueMapper) -> list[PIIMatch]:
    """
    Redact PII within a single paragraph while preserving run-level formatting.

    Strategy:
      1. Concatenate all run texts -> full paragraph text.
      2. Detect PII on the full text (so multi-run names are caught).
      3. Rebuild run texts character-by-character, substituting PII spans.
    """
    # 1. Build full text with per-character run index
    full_text = ""
    run_map: list[int] = []   # run_map[i] = index of run containing char i
    for ridx, run in enumerate(paragraph.runs):
        t = run.text
        full_text += t
        run_map.extend([ridx] * len(t))

    if not full_text.strip():
        return []

    # 2. Fast pre-filter — skip paragraphs that provably contain no PII
    if not _might_contain_pii(full_text):
        return []

    # 3. Detect PII on full text
    matches = _run_detectors(full_text)
    if not matches:
        return []

    # 3. Build replacement text mapping: position -> replacement char sequence
    new_chars: list[str] = list(full_text)
    skip_until = -1
    replacements: dict[int, str] = {}  # start_pos -> replacement string

    for m in matches:
        fake = mapper.get_or_create(m.pii_type, m.text)
        # Mark original chars as "" except the first position which gets the full replacement
        replacements[m.start] = fake
        for i in range(m.start, m.end):
            new_chars[i] = "\x00"   # sentinel: to be removed
        new_chars[m.start] = fake   # first char holds whole replacement

    # 4. Reconstruct per-run text
    run_texts: dict[int, str] = {i: "" for i in range(len(paragraph.runs))}
    for i, ch in enumerate(new_chars):
        if ch == "\x00":
            continue
        ridx = run_map[i] if i < len(run_map) else len(paragraph.runs) - 1
        run_texts[ridx] += ch

    # 5. Apply back to runs
    for ridx, run in enumerate(paragraph.runs):
        run.text = run_texts.get(ridx, "")

    return matches


def _redact_paragraphs(paragraphs, mapper: FakeValueMapper, label: str = "") -> list[PIIMatch]:
    all_matches = []
    para_list = list(paragraphs)
    total = len(para_list)
    skipped = 0
    for idx, para in enumerate(para_list):
        # Log progress every 100 paragraphs so Render logs show activity
        if total > 100 and idx % 100 == 0:
            logger.info(f"  {label}Progress: {idx}/{total} paragraphs ({skipped} skipped by pre-filter)")
        try:
            hits = _redact_paragraph(para, mapper)
            if not hits and para.text.strip() and not _might_contain_pii(para.text):
                skipped += 1
            all_matches.extend(hits)
        except Exception as exc:
            logger.debug(f"Paragraph redaction error: {exc}")
    if total > 0:
        logger.info(f"  {label}Done: {total} paragraphs, {skipped} skipped, {len(all_matches)} PII hits")
    return all_matches


def _redact_table(table, mapper: FakeValueMapper) -> list[PIIMatch]:
    """Iterate all cells in a table and redact each cell's paragraphs."""
    all_matches = []
    for row in table.rows:
        for cell in row.cells:
            hits = _redact_paragraphs(cell.paragraphs, mapper)
            all_matches.extend(hits)
            # Recurse into nested tables
            for nested in cell.tables:
                hits2 = _redact_table(nested, mapper)
                all_matches.extend(hits2)
    return all_matches


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def process_document(
    input_path: str,
    output_path: str,
    mapper: FakeValueMapper,
) -> list[PIIMatch]:
    """
    Read *input_path*, redact all PII, write to *output_path*.
    Returns the flat list of all PIIMatch objects found.
    """
    doc = Document(input_path)
    all_matches: list[PIIMatch] = []

    # --- Body paragraphs ---
    logger.info(f"Redacting body paragraphs ({len(doc.paragraphs)} total)...")
    hits = _redact_paragraphs(doc.paragraphs, mapper, label="Body ")
    all_matches.extend(hits)

    # --- Tables ---
    logger.info("Redacting tables...")
    for table in doc.tables:
        hits = _redact_table(table, mapper)
        all_matches.extend(hits)
    logger.info(f"  Total after tables: {len(all_matches)} PII hits")

    # --- Headers and Footers ---
    logger.info("Redacting headers/footers...")
    for section in doc.sections:
        for hdr_ftr in [section.header, section.footer,
                         section.even_page_header, section.even_page_footer,
                         section.first_page_header, section.first_page_footer]:
            if hdr_ftr is not None:
                try:
                    hits = _redact_paragraphs(hdr_ftr.paragraphs, mapper)
                    all_matches.extend(hits)
                    for table in hdr_ftr.tables:
                        hits = _redact_table(table, mapper)
                        all_matches.extend(hits)
                except Exception:
                    pass  # some sections have no header/footer

    logger.info(f"Total PII detections: {len(all_matches)}")
    doc.save(output_path)
    logger.info(f"Redacted document saved to: {output_path}")
    return all_matches
