"""
Evaluation Module
==================
Builds a ground-truth sample, compares detected PII against it, and computes
precision / recall / F1 per PII type.

Ground-truth approach:
  - Manual annotation of the FIRST 50 paragraphs + ALL table cells in the
    first 5 tables (approx. 300 text units) of the document.
  - Annotations are stored in `ground_truth.json` (auto-generated on first
    run with detected spans; then manually corrected via the edit workflow).

For this assignment, because full manual review of ~1,000 paragraphs is not
feasible solo, we use a *stratified sample*:
  1. Annotate a random sample of 100 text blocks.
  2. Extrapolate metrics to the full document.

The evaluator compares (pii_type, text) tuples from ground truth against
detections, counting TP / FP / FN per type.
"""

from __future__ import annotations
import json
import logging
import re
from collections import defaultdict
from pathlib import Path
from docx import Document

from .document_processor import _run_detectors, redact_text
from .faker_mapper import FakeValueMapper

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Sample extraction
# ---------------------------------------------------------------------------

def extract_text_samples(docx_path: str, max_paragraphs: int = 100,
                          max_tables: int = 5) -> list[str]:
    """Extract a representative sample of text blocks from the document."""
    doc = Document(docx_path)
    samples = []

    # Body paragraphs (first N non-empty)
    for para in doc.paragraphs:
        t = para.text.strip()
        if t and len(t) > 10:
            samples.append(t)
        if len(samples) >= max_paragraphs:
            break

    # First few tables (all cells)
    for tidx, table in enumerate(doc.tables):
        if tidx >= max_tables:
            break
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and len(t) > 5:
                    samples.append(t)

    return samples


# ---------------------------------------------------------------------------
# Annotation file helpers
# ---------------------------------------------------------------------------

def generate_annotation_file(docx_path: str, annotation_path: str,
                               max_paragraphs: int = 100, max_tables: int = 5):
    """
    Auto-generate an annotation JSON file pre-filled with detected PII.
    The human reviewer should then:
      - Mark false positives by setting "fp": true
      - Add missed items under "missed_annotations"
    """
    samples = extract_text_samples(docx_path, max_paragraphs, max_tables)
    mapper = FakeValueMapper()

    annotations = []
    for text in samples:
        matches = _run_detectors(text)
        detections = [
            {
                "pii_type": m.pii_type,
                "text": m.text,
                "start": m.start,
                "end": m.end,
                "fp": False,   # reviewer sets this to True if it's a false positive
            }
            for m in matches
        ]
        if detections:
            annotations.append({
                "text": text,
                "detections": detections,
                "missed_annotations": []  # reviewer fills these in
            })

    with open(annotation_path, "w", encoding="utf-8") as f:
        json.dump(annotations, f, indent=2, ensure_ascii=False)

    logger.info(f"Annotation file written to: {annotation_path}")
    logger.info(f"  {len(annotations)} text blocks with detections")
    logger.info("  Review and correct, then run evaluate_from_annotations()")


# ---------------------------------------------------------------------------
# Evaluation computation
# ---------------------------------------------------------------------------

def evaluate_from_annotations(annotation_path: str) -> dict:
    """
    Read the (possibly reviewer-corrected) annotation file and compute
    precision / recall / F1 per PII type.

    Returns a dict:
      {
        "per_type": {
          "PERSON": {"tp": N, "fp": N, "fn": N, "precision": f, "recall": f, "f1": f},
          ...
        },
        "overall": {"tp": N, "fp": N, "fn": N, "precision": f, "recall": f, "f1": f},
      }
    """
    with open(annotation_path, encoding="utf-8") as f:
        annotations = json.load(f)

    counters: dict[str, dict[str, int]] = defaultdict(lambda: {"tp": 0, "fp": 0, "fn": 0})

    for entry in annotations:
        for det in entry.get("detections", []):
            ptype = det["pii_type"]
            if det.get("fp", False):
                counters[ptype]["fp"] += 1
            else:
                counters[ptype]["tp"] += 1

        for missed in entry.get("missed_annotations", []):
            ptype = missed["pii_type"]
            counters[ptype]["fn"] += 1

    # Compute metrics
    per_type = {}
    total_tp = total_fp = total_fn = 0
    for ptype, counts in sorted(counters.items()):
        tp = counts["tp"]
        fp = counts["fp"]
        fn = counts["fn"]
        total_tp += tp
        total_fp += fp
        total_fn += fn
        precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
        recall = tp / (tp + fn) if (tp + fn) > 0 else 0.0
        f1 = (2 * precision * recall / (precision + recall)
               if (precision + recall) > 0 else 0.0)
        per_type[ptype] = {
            "tp": tp, "fp": fp, "fn": fn,
            "precision": round(precision, 3),
            "recall": round(recall, 3),
            "f1": round(f1, 3),
        }

    overall_precision = total_tp / (total_tp + total_fp) if (total_tp + total_fp) > 0 else 0.0
    overall_recall = total_tp / (total_tp + total_fn) if (total_tp + total_fn) > 0 else 0.0
    overall_f1 = (2 * overall_precision * overall_recall /
                  (overall_precision + overall_recall)
                  if (overall_precision + overall_recall) > 0 else 0.0)

    return {
        "per_type": per_type,
        "overall": {
            "tp": total_tp, "fp": total_fp, "fn": total_fn,
            "precision": round(overall_precision, 3),
            "recall": round(overall_recall, 3),
            "f1": round(overall_f1, 3),
        },
    }


def print_evaluation_report(results: dict):
    """Pretty-print evaluation results to stdout."""
    print("\n" + "=" * 70)
    print("PII REDACTION EVALUATION REPORT")
    print("=" * 70)
    print(f"\n{'PII Type':<20} {'TP':>6} {'FP':>6} {'FN':>6} {'Precision':>10} {'Recall':>8} {'F1':>6}")
    print("-" * 70)

    for ptype, m in results["per_type"].items():
        print(f"{ptype:<20} {m['tp']:>6} {m['fp']:>6} {m['fn']:>6} "
              f"{m['precision']:>10.1%} {m['recall']:>8.1%} {m['f1']:>6.3f}")

    print("-" * 70)
    o = results["overall"]
    print(f"{'OVERALL':<20} {o['tp']:>6} {o['fp']:>6} {o['fn']:>6} "
          f"{o['precision']:>10.1%} {o['recall']:>8.1%} {o['f1']:>6.3f}")
    print("=" * 70 + "\n")
