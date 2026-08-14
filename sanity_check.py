"""
Quick sanity check: verify the redacted .docx has actual substitutions
and that known PII values from the original are NOT present in the redacted file.
"""
from docx import Document

orig_path = "Red Herring Prospectus.docx"
redacted_path = "Red_Herring_Prospectus_Redacted.docx"

# Known PII values that MUST be absent from the redacted file
MUST_BE_ABSENT = [
    "Sarthak Malvadkar",
    "cs.connect@kshinternational.com",
    "sarthak.malvadkar@kshinterantional.com",
    "4505 3237",          # part of the landline that should be redacted
    "U28129PN1979PLC141032",
]

def extract_all_text(path):
    doc = Document(path)
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)

print("Extracting text from original...")
orig_text = extract_all_text(orig_path)
print("Extracting text from redacted...")
redacted_text = extract_all_text(redacted_path)

print(f"\nOriginal length:  {len(orig_text):,} chars")
print(f"Redacted length:  {len(redacted_text):,} chars\n")

print("=== PII Presence Check ===")
all_passed = True
for pii in MUST_BE_ABSENT:
    in_orig = pii.lower() in orig_text.lower()
    in_redacted = pii.lower() in redacted_text.lower()
    status = "[PASS]" if in_orig and not in_redacted else ("[NOT IN ORIG]" if not in_orig else "[FAIL -- still present!]")
    print(f"  {status}  {pii!r}")
    if in_orig and in_redacted:
        all_passed = False

print(f"\n{'[OK] All PII successfully removed' if all_passed else '[ERROR] Some PII still present -- check logs'}")
